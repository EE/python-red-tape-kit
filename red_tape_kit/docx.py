from unittest.mock import patch
from zipfile import ZipInfo

import docx
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from .doc_ast import DefinitionList, Image, Paragraph, Section, Sequence, Table, UnorderedList


class DOCXRenderer:
    def __init__(self, document):
        self.document = document
        self.docx = docx.Document()

        self.docx.styles['Header'].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        self.docx.styles['Header'].paragraph_format.space_after = docx.shared.Pt(6)
        self.docx.styles['Footer'].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        self.docx.styles['Footer'].paragraph_format.space_before = docx.shared.Pt(6)

        self.set_meta()
        self.add_cover()
        self.add_body()

    def set_meta(self):
        self.docx.core_properties.language = self.document.language_code
        self.docx.core_properties.title = self.document.title.plain_string
        self.docx.core_properties.subject = self.document.subject.plain_string
        self.docx.core_properties.author = self.document.author.plain_string
        self.docx.core_properties.creator = self.document.creator.plain_string
        self.docx.core_properties.created = self.document.creation_date

    def add_cover(self):
        assert len(self.docx.sections) == 1
        section = self.docx.sections[0]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.top_margin = docx.shared.Inches(4)
        section.left_margin = docx.shared.Inches(3)
        self.docx.add_heading(self.document.title.plain_string, 0)
        self.docx.add_paragraph(self.document.subject.plain_string)
        self.docx.add_paragraph(
            f'Autor: {self.document.author.plain_string}\n'
            f'{self.document.creation_place_and_date}'
        )

    def add_body(self):
        body_section = self.docx.add_section()
        body_section.top_margin = docx.shared.Inches(1)
        body_section.left_margin = docx.shared.Inches(1.25)

        # header config
        body_section.header.is_linked_to_previous = False
        p = body_section.header.add_paragraph(
            f'{self.document.title.plain_string}\n{self.document.subject.plain_string}',
            style='Header',
        )
        # TODO: header border should be ideally defined in the "Header" style
        self.set_hr(p, location='bottom')

        # footer config
        body_section.footer.is_linked_to_previous = False
        p = body_section.footer.add_paragraph(
            f'{self.document.author.plain_string}\n{self.document.creation_place_and_date}',
            style='Footer',
        )
        # TODO: footer border should be ideally defined in the "Footer" style
        self.set_hr(p, location='top')

        self.add_element(self.document.body, section_level=1)

    def add_element(self, element, section_level=None, list_level=None, first_in_list_item=False):
        if isinstance(element, Section):
            self.add_section(element, section_level)
        elif isinstance(element, Paragraph):
            self.add_paragraph(element, list_level, first_in_list_item)
        elif isinstance(element, Table):
            self.add_table(element)
        elif isinstance(element, UnorderedList):
            self.add_unordered_list(element, list_level)
        elif isinstance(element, DefinitionList):
            self.add_definition_list(element)
        elif isinstance(element, Sequence):
            self.add_sequence(element, section_level, list_level, first_in_list_item)
        elif isinstance(element, Image):
            self.add_image(element)
        else:
            raise ValueError(f'Unknown element type {element}')

    def add_sequence(self, sequence, section_level, list_level, first_in_list_item):
        for sub_element in sequence.items:
            self.add_element(sub_element, section_level, list_level, first_in_list_item)
            first_in_list_item = False

    def add_section(self, section, section_level):
        self.docx.add_heading(section.title.plain_string, section_level)
        self.add_element(section.body, section_level + 1)

    def add_paragraph(self, paragraph, list_level, first_in_list_item):
        if not first_in_list_item and list_level is None:
            style = 'Body Text'
        elif not first_in_list_item and list_level == 1:
            style = 'List Continue'
        elif not first_in_list_item and list_level > 1:
            style = f'List Continue {list_level}'
        elif first_in_list_item and list_level == 1:
            style = 'List Bullet'
        elif first_in_list_item and list_level > 1:
            style = f'List Bullet {list_level}'
        else:
            raise ValueError(
                'Invalid paragraph config: '
                f'list_level={list_level}, first_in_list_item={first_in_list_item}'
            )
        self.docx.add_paragraph(paragraph.text.plain_string, style=style)

    def add_table(self, table_data):
        table = self.docx.add_table(rows=1, cols=len(table_data.headings))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, heading in enumerate(table_data.headings):
            hdr_cells[i].text = heading.plain_string
        for data_row in table_data.rows:
            row_cells = table.add_row().cells
            for i, data_cell in enumerate(data_row):
                row_cells[i].text = data_cell.plain_string

    def add_unordered_list(self, unordered_list, list_level):
        if list_level is None:
            list_level = 0
        for item in unordered_list.items:
            self.add_element(item, list_level=list_level + 1, first_in_list_item=True)

    def add_definition_list(self, definition_list):
        list_items = []
        for key, value in definition_list.items.items():
            list_items.append(Sequence([
                Paragraph(key),
                value,
            ]))
        self.add_element(UnorderedList(list_items), list_level=0)

    def add_image(self, image):
        self.docx.add_picture(image.image_io, width=docx.shared.Inches(6))
        self.docx.add_paragraph(image.caption.plain_string)

    def set_hr(self, paragraph, location='top'):
        """set bottom border on a paragraph
        taken from https://github.com/python-openxml/python-docx/issues/105#issuecomment-442786431
        """
        assert location in ('top', 'bottom')
        p = paragraph._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(
            pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange',
        )
        bottom = OxmlElement(f'w:{location}')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def render(self, file):
        # python-docx package sets current time when packaging a document (zipfile)
        # We fix the time to achive reproducible output.
        # See:
        # - https://github.com/python-openxml/python-docx/issues/811
        # - https://github.com/python-openxml/python-docx/issues/1042
        def fixed_write(self, pack_uri, blob):
            zinfo = ZipInfo(filename=pack_uri.membername)
            self._zipf.writestr(zinfo, blob)
        with patch('docx.opc.phys_pkg._ZipPkgWriter.write', fixed_write):
            self.docx.save(file)
